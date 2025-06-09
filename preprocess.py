#preprocess.py

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from fastapi import UploadFile
from dotenv import load_dotenv
from io import BytesIO
import os
import faiss
import pickle
import numpy as np
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
import sys
import sqlite3
from langchain_community.vectorstores import FAISS

# Fix for FAISS SQLite dependency in some environments
sys.modules["sqlite3"] = sqlite3

load_dotenv()

# Preprocess uploaded files + scraped data into text chunks
async def preprocess_text(files: list[UploadFile], size, overlap ,scraped_data): #scraped_data
    import time
    
    paragraphs = []
    from io import BytesIO
    # Step 1: Process each file
    for file in files:
        if file.filename.endswith(".pdf"):
            contents = await file.read()
            file_object = BytesIO(contents)
            reader = PdfReader(file_object)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    cleaned_text = ' '.join(page_text.split())
                    paragraphs.append(cleaned_text)
        elif file.filename.endswith(".docx"):
            
            file_content = await file.read() 
            docx_stream = BytesIO(file_content) 
            docx = DocxDocument(docx_stream)
            full_text = ""
            for para in docx.paragraphs:
                if para.text.strip():
                    full_text += para.text.strip() + "\n\n"
            paragraphs.append(full_text)

    print(f"📄 Total paragraphs after split: {len(paragraphs)}")
    print("🧩 First 5 extracted paragraphs:")
    for i, p in enumerate(paragraphs[:5]):
        print(f"{i+1}. {p[:100]}...")

    if scraped_data:
        if isinstance(scraped_data, str):
            paragraphs.extend(scraped_data.split("\n\n"))  # Break on double newline
        elif isinstance(scraped_data, list):
            for item in scraped_data:
                if isinstance(item, dict) and 'full_text' in item:
                    chunks = [chunk.strip() for chunk in item['full_text'].split("\n\n") if chunk.strip()]
                    paragraphs.extend(chunks)


    paragraphs = [para.strip() for para in paragraphs if para.strip()]
    # print(paragraphs)

    docs = [LangchainDocument(page_content=para) for para in paragraphs]

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=size, chunk_overlap=overlap)
    text_chunks = text_splitter.split_documents(docs)
    # print(text_chunks)
    return text_chunks

# Main entrypoint to support multiple vector DBs
async def preprocess_vectordbs(
    doc_files, embedding_model_name, chunk_size, chunk_overlap, scraped_data,
    selected_vectordb, persist_directory=None
):
    print(f"[INFO] Preprocessing for vector DB: {selected_vectordb}")

    # ✅ Check: if both PDF and scraped data are missing, raise error
    if not doc_files and not scraped_data:
        raise ValueError("No documents or scraped content to process.")

    texts = await preprocess_text(doc_files, chunk_size, chunk_overlap, scraped_data)
    print(f"[DEBUG] Number of documents/chunks: {len(texts)}")
    for i, doc in enumerate(texts[:10]):  # Check the first 10 chunks
        print(f"[DEBUG] Chunk {i+1} content preview: {repr(doc.page_content[:100])}")
        if not doc.page_content.strip():
            print(f"[WARNING] Chunk {i+1} is empty or whitespace!")

    print(f"[INFO] Initializing embedding model: {embedding_model_name}")
    embedding_model = SentenceTransformerEmbeddings(model_name=embedding_model_name)
    embedding_vector = embedding_model.embed_query("test")
    # print("Embedding dimension:", len(embedding_vector))

    if selected_vectordb == "FAISS":
        print("[INFO] Building FAISS vectorstore...")
        print(f"[INFO] Total document chunks: {len(texts)}")
        for i, doc in enumerate(texts):
            if not doc.page_content.strip():
                print(f"[⚠️] Empty content at chunk {i}")
            else:
                print(f"[✅] Chunk {i} preview: {doc.page_content[:80]}...")

        # ✅ Rebuild from scratch
        vectorstore = FAISS.from_documents(texts, embedding_model)

        if persist_directory:
            print(f"[INFO] Saving FAISS index manually to: {persist_directory}")
            os.makedirs(persist_directory, exist_ok=True)

            # Save FAISS index
            faiss.write_index(vectorstore.index, os.path.join(persist_directory, "index.faiss"))

            # Save docstore and mapping
            with open(os.path.join(persist_directory, "index.pkl"), "wb") as f:
                pickle.dump({
                    "docstore": vectorstore.docstore,
                    "index_to_docstore_id": vectorstore.index_to_docstore_id
                }, f)

            # ✅ DEBUG: Confirm docstore validity
            print("✅ Number of documents in docstore:", len(vectorstore.docstore._dict))
            print("✅ Sample docstore entries:")
            for i, (k, v) in enumerate(vectorstore.docstore._dict.items()):
                print(f"  {i+1}. Key: {k} → Content: {v.page_content[:100] if v else 'None'}")
                if i >= 4: break  # print first 5 only

            retriever = vectorstore.as_retriever()
            print("[INFO] FAISS vectorstore ready.")
            return (
                vectorstore.index,
                vectorstore.docstore,
                vectorstore.index_to_docstore_id,
                vectorstore,
                retriever,
                embedding_model,
                None,  # pinecone_index_name
                None,  # vs
                None   # qdrant_client
            )

    # ❌ Unsupported DB case
    raise ValueError(f"[ERROR] Unsupported vector DB selected: {selected_vectordb}")


